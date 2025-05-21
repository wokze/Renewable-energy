import pulp
import random
import numpy as np  # Add this import statement
import matplotlib.pyplot as plt  # Add this import

# Wind turbine parameters
Pr = 20  # Rated power (kW)
vci = 3  # Cut-in wind speed (m/s)
vco = 25  # Cut-out wind speed (m/s)
vN = 12  # Nominal wind speed (m/s)

# Constants for the power curve
A = 0.3
B = 0.4
C = 0.3

# Generate random wind speeds for 24 hours
np.random.seed(42)
wind_speeds = np.random.uniform(0, 30, 24)

# Create optimization problem
prob = pulp.LpProblem("Wind_Turbine_Optimization", pulp.LpMaximize)

# Decision variables for power output at each hour
power_vars = [pulp.LpVariable(f"power_{t}", 0, Pr) for t in range(24)]

# Function to calculate power output based on wind speed
def calculate_power(v):
    if v < vci or v > vco:
        return 0
    elif vci <= v <= vN:
        # Linear interpolation between cut-in speed and nominal speed
        return Pr * ((v - vci) / (vN - vci))
    elif vN < v <= vco:
        return Pr
    
# Objective function: maximize total power output
prob += pulp.lpSum(power_vars)

# Constraints
for t in range(24):
    v = wind_speeds[t]
    max_power = calculate_power(v)
    prob += power_vars[t] <= max_power
    prob += power_vars[t] >= 0

# Solve the problem
prob.solve()

# PV System Parameters
Pn = 15  # Nominal power of PV unit (kW)
beta = 0.9  # Solar irradiance factor
k = -0.005  # Temperature coefficient
Ts = 25  # Standard temperature (°C)

# Generate solar irradiance data (bell curve for 24 hours)
hour_range = np.linspace(0, 23, 24)
solar_irradiance = np.zeros(24)
for i in range(24):
    if 6 <= i <= 18:  # Daylight hours (6 AM to 6 PM)
        solar_irradiance[i] = np.exp(-((i-12)**2)/(2*3**2))  # Bell curve centered at noon
solar_irradiance = solar_irradiance * 1000  # Scale to W/m²

# Temperature variation during the day
temperature = 20 + 10 * solar_irradiance/1000  # Temperature varies between 20-30°C

# Function to calculate PV power output
def calculate_pv_power(hour):
    G = solar_irradiance[hour]
    Tc = temperature[hour]
    if G > 0:
        return Pn * beta * (G/1000) * (1 + k * (Tc - Ts))
    return 0

# Create separate PV optimization problem
pv_prob = pulp.LpProblem("PV_Optimization", pulp.LpMaximize)

# Decision variables for PV power output
pv_power_vars = [pulp.LpVariable(f"pv_power_{t}", 0, Pn) for t in range(24)]

# Objective function for PV: maximize total PV power output
pv_prob += pulp.lpSum(pv_power_vars)

# PV Constraints
for t in range(24):
    max_pv_power = calculate_pv_power(t)
    pv_prob += pv_power_vars[t] <= max_pv_power
    pv_prob += pv_power_vars[t] >= 0

# Solve PV optimization
pv_prob.solve()

# Battery Parameters
Ee = 50  # Battery capacity (kWh)
eta_s = 0.95  # Charging efficiency
eta_d = 0.95  # Discharging efficiency
SoC_min = 0.1 * Ee  # Minimum state of charge (10% of capacity)
SoC_max = 0.9 * Ee  # Maximum state of charge (90% of capacity)
P_max = 10  # Maximum charging/discharging power (kW)

# Time-of-use electricity prices (€/kWh)
electricity_prices = np.zeros(24)
electricity_prices[8:12] = 0.25  # Peak morning
electricity_prices[18:22] = 0.25 # Peak evening
electricity_prices[6:8] = 0.15   # Shoulder morning
electricity_prices[12:18] = 0.15 # Shoulder afternoon
electricity_prices[22:23] = 0.15 # Shoulder evening
electricity_prices[23:] = 0.10   # Off-peak night
electricity_prices[:6] = 0.10    # Off-peak early morning

# Load consumption profile (kW) - typical daily pattern
load_demand = np.zeros(24)
# Morning peak (6-9)
load_demand[6:9] = 8
# Day consumption (9-17)
load_demand[9:17] = 5
# Evening peak (17-22)
load_demand[17:22] = 10
# Night consumption (22-6)
load_demand[22:] = 3
load_demand[:6] = 3

# Calculate net power before battery (generation - consumption)
net_power = np.zeros(24)
for t in range(24):
    generation = power_vars[t].value() + pv_power_vars[t].value()
    net_power[t] = generation - load_demand[t]

# Create combined battery optimization problem
bat_prob = pulp.LpProblem("Battery_Economic_Optimization", pulp.LpMaximize)

# Add binary variables to control charging/discharging states
charge_state = [pulp.LpVariable(f"charge_state_{t}", 0, 1, pulp.LpBinary) for t in range(24)]

# Decision variables for battery
soc_vars = [pulp.LpVariable(f"soc_{t}", SoC_min, SoC_max) for t in range(25)]  # Battery state of charge
charge_vars = [pulp.LpVariable(f"charge_{t}", 0, P_max) for t in range(24)]     # When to charge
discharge_vars = [pulp.LpVariable(f"discharge_{t}", 0, P_max) for t in range(24)] # When to discharge
grid_import = [pulp.LpVariable(f"grid_import_{t}", 0, None) for t in range(24)]
grid_export = [pulp.LpVariable(f"grid_export_{t}", 0, None) for t in range(24)]

# Add variables for tracking charge/discharge rate changes
charge_change = [pulp.LpVariable(f"charge_change_{t}", -P_max, P_max) for t in range(23)]
discharge_change = [pulp.LpVariable(f"discharge_change_{t}", -P_max, P_max) for t in range(23)]

# Add auxiliary variables for absolute values
charge_change_pos = [pulp.LpVariable(f"charge_change_pos_{t}", 0, P_max) for t in range(23)]
charge_change_neg = [pulp.LpVariable(f"charge_change_neg_{t}", 0, P_max) for t in range(23)]
discharge_change_pos = [pulp.LpVariable(f"discharge_change_pos_{t}", 0, P_max) for t in range(23)]
discharge_change_neg = [pulp.LpVariable(f"discharge_change_neg_{t}", 0, P_max) for t in range(23)]

# Objective function: maximize profit
bat_prob += (pulp.lpSum([grid_export[t] * electricity_prices[t] for t in range(24)]) -  # Revenue from selling
             pulp.lpSum([grid_import[t] * electricity_prices[t] for t in range(24)]))   # Cost of buying

# Battery constraints
for t in range(24):
    # Prevent simultaneous charging and discharging
    bat_prob += charge_vars[t] <= P_max * charge_state[t]
    bat_prob += discharge_vars[t] <= P_max * (1 - charge_state[t])
    
    # Sequential transition constraints
    if t > 0:
        # If we were charging in previous hour, ensure charging reaches zero before discharging starts
        bat_prob += discharge_vars[t] <= P_max * (1 - charge_state[t-1])  # Can only discharge if not charging in previous hour
        bat_prob += charge_vars[t] <= P_max * charge_state[t-1]  # Can only charge if charging in previous hour
        
        # Rate of change limits
        bat_prob += charge_vars[t] <= charge_vars[t-1]  # Charging can only decrease
        bat_prob += discharge_vars[t] <= P_max * (1 - charge_state[t-1])  # Discharge limited by previous charging state

    # SoC evolution constraint
    bat_prob += Ee * eta_d * soc_vars[t+1] == Ee * eta_d * soc_vars[t] + \
                (charge_vars[t] * eta_s * eta_d) - discharge_vars[t]
    
    # Power balance constraint
    bat_prob += net_power[t] + discharge_vars[t] - charge_vars[t] + grid_import[t] - grid_export[t] == 0
    
    # Prevent extreme charging/discharging changes
    if t < 23:
        bat_prob += charge_vars[t+1] - charge_vars[t] == charge_change[t]
        bat_prob += discharge_vars[t+1] - discharge_vars[t] == discharge_change[t]
        # Define absolute value constraints
        bat_prob += charge_change[t] == charge_change_pos[t] - charge_change_neg[t]
        bat_prob += discharge_change[t] == discharge_change_pos[t] - discharge_change_neg[t]
        # Limit rate of change to 30% of P_max per hour
        bat_prob += charge_change[t] <= 0.3 * P_max
        bat_prob += charge_change[t] >= -0.3 * P_max
        bat_prob += discharge_change[t] <= 0.3 * P_max
        bat_prob += discharge_change[t] >= -0.3 * P_max

# Modify objective function to use auxiliary variables instead of abs()
smoothing_factor = 0.1  # Penalty factor for rapid changes
bat_prob += (pulp.lpSum([grid_export[t] * electricity_prices[t] for t in range(24)]) - 
             pulp.lpSum([grid_import[t] * electricity_prices[t] for t in range(24)]) -
             smoothing_factor * pulp.lpSum([(charge_change_pos[t] + charge_change_neg[t] + 
                                           discharge_change_pos[t] + discharge_change_neg[t]) 
                                          for t in range(23)]))

# Initial and final SoC constraints
bat_prob += soc_vars[0] == 0.5 * Ee
bat_prob += soc_vars[24] >= 0.5 * Ee

# Solve battery optimization
bat_prob.solve()

# Print both wind and PV results
print("Wind speeds and optimized power output:")
for t in range(24):
    print(f"Hour {t}: Wind Speed = {wind_speeds[t]:.2f} m/s, Power = {power_vars[t].value():.2f} kW")

print("\nPV optimized power output:")
for t in range(24):
    print(f"Hour {t}: Solar Irradiance = {solar_irradiance[t]:.2f} W/m², Power = {pv_power_vars[t].value():.2f} kW")

# Print battery results with prices
print("\nBattery Operation Results:")
for t in range(24):
    print(f"Hour {t}:")
    print(f"  Price = {electricity_prices[t]:.2f} €/kWh")
    print(f"  SoC = {soc_vars[t].value():.2f} kWh")
    print(f"  Charging = {charge_vars[t].value():.2f} kW")
    print(f"  Discharging = {discharge_vars[t].value():.2f} kW")

# Print comprehensive results
print("\nHourly Power Flow Analysis:")
for t in range(24):
    print(f"\nHour {t}:")
    print(f"  Load Demand: {load_demand[t]:.2f} kW")
    print(f"  Generation: {power_vars[t].value() + pv_power_vars[t].value():.2f} kW")
    print(f"  Battery Charging: {charge_vars[t].value():.2f} kW")
    print(f"  Battery Discharging: {discharge_vars[t].value():.2f} kW")
    print(f"  Grid Import: {grid_import[t].value():.2f} kW")
    print(f"  Grid Export: {grid_export[t].value():.2f} kW")
    print(f"  Electricity Price: {electricity_prices[t]:.2f} €/kWh")
    print(f"  Battery SoC: {soc_vars[t].value():.2f} kWh")

# Calculate total system cost F
print("\nTotal System Cost Analysis:")
total_cost = 0
hours = list(range(24))
wind_costs = []
pv_costs = []
battery_costs = []
total_hourly_costs = []

for t in range(24):
    # Wind turbine component
    wind_cost = power_vars[t].value() * electricity_prices[t]
    wind_costs.append(wind_cost)
    
    # PV component
    pv_cost = pv_power_vars[t].value() * electricity_prices[t]
    pv_costs.append(pv_cost)
    
    # Battery component (charging and discharging)
    battery_cost = (charge_vars[t].value() + discharge_vars[t].value()) * electricity_prices[t]
    battery_costs.append(battery_cost)
    
    # Sum all components for this hour
    hourly_cost = wind_cost + pv_cost + battery_cost
    total_hourly_costs.append(hourly_cost)
    total_cost += hourly_cost
    
    print(f"\nHour {t} Cost Breakdown:")
    print(f"  Wind Cost: {wind_cost:.2f} €")
    print(f"  PV Cost: {pv_cost:.2f} €")
    print(f"  Battery Operation Cost: {battery_cost:.2f} €")
    print(f"  Total Hourly Cost: {hourly_cost:.2f} €")

print(f"\nTotal System Cost F = {total_cost:.2f} €")

# Create figure for component analysis
plt.figure(figsize=(15, 10))

# Wind Turbine Power Output
plt.subplot(2, 2, 1)
plt.plot(hours, [power_vars[t].value() for t in range(24)], 'b-', label='Wind Power')
plt.fill_between(hours, [power_vars[t].value() for t in range(24)], alpha=0.3)
plt.xlabel('Hour')
plt.ylabel('Power (kW)')
plt.title('Wind Turbine Power Output')
plt.grid(True, alpha=0.3)
plt.legend()

# PV System Power Output
plt.subplot(2, 2, 2)
plt.plot(hours, [pv_power_vars[t].value() for t in range(24)], 'y-', label='PV Power')
plt.fill_between(hours, [pv_power_vars[t].value() for t in range(24)], color='yellow', alpha=0.3)
plt.xlabel('Hour')
plt.ylabel('Power (kW)')
plt.title('PV System Power Output')
plt.grid(True, alpha=0.3)
plt.legend()

# Battery Charging/Discharging
plt.subplot(2, 2, 3)
plt.plot(hours, [charge_vars[t].value() for t in range(24)], 'g-', label='Charging')
plt.plot(hours, [-discharge_vars[t].value() for t in range(24)], 'r-', label='Discharging')
plt.fill_between(hours, [charge_vars[t].value() for t in range(24)], color='green', alpha=0.3)
plt.fill_between(hours, [-discharge_vars[t].value() for t in range(24)], color='red', alpha=0.3)
plt.xlabel('Hour')
plt.ylabel('Power (kW)')
plt.title('Battery Operation')
plt.grid(True, alpha=0.3)
plt.legend()

# Battery State of Charge
plt.subplot(2, 2, 4)
plt.plot(hours, [soc_vars[t].value() for t in range(24)], 'b-', label='State of Charge')
plt.fill_between(hours, [soc_vars[t].value() for t in range(24)], alpha=0.3)
plt.xlabel('Hour')
plt.ylabel('Energy (kWh)')
plt.title('Battery State of Charge')
plt.grid(True, alpha=0.3)
plt.legend()

plt.tight_layout()
plt.show()

# Create figure with two subplots side by side
plt.figure(figsize=(15, 6))

# First subplot - Variable price scenario
plt.subplot(1, 2, 1)
plt.bar(hours, wind_costs, label='Wind Cost', color='blue', alpha=0.5)
plt.bar(hours, pv_costs, bottom=wind_costs, label='PV Cost', color='yellow', alpha=0.5)
plt.bar(hours, battery_costs, bottom=[sum(x) for x in zip(wind_costs, pv_costs)], 
        label='Battery Cost', color='green', alpha=0.5)

plt.xlabel('Hour')
plt.ylabel('Cost (€)')
plt.title('System Cost - Variable Price')
plt.legend()
plt.grid(True, alpha=0.3)
plt.text(0.02, 0.95, f'Total System Cost: {sum(total_hourly_costs):.2f} €', 
         transform=plt.gca().transAxes, bbox=dict(facecolor='white', alpha=0.8))

# Calculate costs for fixed price scenario
fixed_wind_costs = [power_vars[t].value() * 0.1268 for t in range(24)]
fixed_pv_costs = [pv_power_vars[t].value() * 0.1268 for t in range(24)]
fixed_battery_costs = [(charge_vars[t].value() - discharge_vars[t].value()) * 0.1268 for t in range(24)]
fixed_total_costs = [sum(x) for x in zip(fixed_wind_costs, fixed_pv_costs, fixed_battery_costs)]

# Second subplot - Fixed price scenario
plt.subplot(1, 2, 2)
plt.bar(hours, fixed_wind_costs, label='Wind Cost', color='blue', alpha=0.5)
plt.bar(hours, fixed_pv_costs, bottom=fixed_wind_costs, label='PV Cost', color='yellow', alpha=0.5)
plt.bar(hours, fixed_battery_costs, bottom=[sum(x) for x in zip(fixed_wind_costs, fixed_pv_costs)], 
        label='Battery Cost', color='green', alpha=0.5)

plt.xlabel('Hour')
plt.ylabel('Cost (€)')
plt.title('System Cost - Fixed Price (0.1268 €/kWh)')
plt.legend()
plt.grid(True, alpha=0.3)
plt.text(0.02, 0.95, f'Total System Cost: {sum(fixed_total_costs):.2f} €', 
         transform=plt.gca().transAxes, bbox=dict(facecolor='white', alpha=0.8))

plt.tight_layout()
plt.show()

# Add final equation explanation and calculation
print("\nFinal Cost Equation:")
print("F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t - Pdischarge,t)) * price(t)]")
print("\nCalculation Breakdown:")
for t in range(24):
    wind_term = power_vars[t].value()
    pv_term = pv_power_vars[t].value()
    battery_term = charge_vars[t].value() + discharge_vars[t].value()  # Changed to addition
    price = electricity_prices[t]
    hour_result = (wind_term + pv_term + battery_term) * price
    
    print(f"\nHour {t}:")
    print(f"  [({wind_term:.2f} + {pv_term:.2f} + ({battery_term:.2f})) * {price:.2f}]")
    print(f"  = [{wind_term + pv_term + battery_term:.2f} * {price:.2f}]")
    print(f"  = {hour_result:.2f} €")

print(f"\nFinal Equation Result F = {total_cost:.2f} €")

# Replace the time-of-use electricity prices with fixed price
electricity_prices = np.full(24, 0.1268)  # Fixed price of 0.1268 €/kWh

# Calculate total system cost F with new fixed price
print("\nTotal System Cost Analysis with Fixed Price (0.1268 €/kWh):")
total_cost = 0
for t in range(24):
    # Wind turbine component
    wind_term = power_vars[t].value()
    pv_term = pv_power_vars[t].value()
    battery_term = charge_vars[t].value() - discharge_vars[t].value()
    
    # Calculate hourly cost with fixed price
    hourly_cost = (wind_term + pv_term + battery_term) * 0.1268
    total_cost += hourly_cost
    
    print(f"\nHour {t}:")
    print(f"  [({wind_term:.2f} + {pv_term:.2f} + ({battery_term:.2f})) * 0.1268]")
    print(f"  = [{wind_term + pv_term + battery_term:.2f} * 0.1268]")
    print(f"  = {hourly_cost:.2f} €")

print(f"\nFinal Cost with Fixed Price = {total_cost:.2f} €")

# Calculate final cost equation
final_hourly_costs = []
final_total_cost = 0
wind_component = []
pv_component = []
battery_component = []

for t in range(24):
    wind_term = power_vars[t].value() * electricity_prices[t]
    pv_term = pv_power_vars[t].value() * electricity_prices[t]
    battery_term = (charge_vars[t].value() - discharge_vars[t].value()) * electricity_prices[t]
    
    wind_component.append(wind_term)
    pv_component.append(pv_term)
    battery_component.append(battery_term)
    
    hourly_cost = wind_term + pv_term + battery_term
    final_hourly_costs.append(hourly_cost)
    final_total_cost += hourly_cost

# Create visualization for final cost equation
plt.figure(figsize=(12, 6))
plt.bar(hours, wind_component, label='Wind Component', color='blue', alpha=0.5)
plt.bar(hours, pv_component, bottom=wind_component, label='PV Component', color='yellow', alpha=0.5)
plt.bar(hours, battery_component, 
        bottom=[sum(x) for x in zip(wind_component, pv_component)],
        label='Battery Component (Charge - Discharge)', color='green', alpha=0.5)

plt.xlabel('Hour')
plt.ylabel('Cost (€)')
plt.title('Final Cost Equation Components')
plt.legend()
plt.grid(True, alpha=0.3)

# Add total cost annotation
plt.text(0.02, 0.95, f'Total Cost F = {final_total_cost:.2f} €', 
         transform=plt.gca().transAxes, 
         bbox=dict(facecolor='white', alpha=0.8))

plt.tight_layout()
plt.show()

print("\nFinal Cost Equation Breakdown:")
print("F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t - Pdischarge,t)) * price(t)]")
print(f"Total F = {final_total_cost:.2f} €")

# Reset electricity prices to time-of-use prices
electricity_prices = np.zeros(24)
electricity_prices[8:12] = 0.25  # Peak morning
electricity_prices[18:22] = 0.25 # Peak evening
electricity_prices[6:8] = 0.15   # Shoulder morning
electricity_prices[12:18] = 0.15 # Shoulder afternoon
electricity_prices[22:23] = 0.15 # Shoulder evening
electricity_prices[23:] = 0.10   # Off-peak night
electricity_prices[:6] = 0.10    # Off-peak early morning

# Calculate final cost equation with time-of-use prices
final_hourly_costs = []
final_total_cost = 0
wind_component = []
pv_component = []
battery_component = []

for t in range(24):
    wind_term = power_vars[t].value() * electricity_prices[t]
    pv_term = pv_power_vars[t].value() * electricity_prices[t]
    battery_term = (charge_vars[t].value() + discharge_vars[t].value()) * electricity_prices[t]
    
    wind_component.append(wind_term)
    pv_component.append(pv_term)
    battery_component.append(battery_term)
    
    hourly_cost = wind_term + pv_term + battery_term
    final_hourly_costs.append(hourly_cost)
    final_total_cost += hourly_cost

# Create visualization with actual prices noted
plt.figure(figsize=(12, 6))
plt.bar(hours, wind_component, label='Wind Component', color='blue', alpha=0.5)
plt.bar(hours, pv_component, bottom=wind_component, label='PV Component', color='yellow', alpha=0.5)
plt.bar(hours, battery_component, 
        bottom=[sum(x) for x in zip(wind_component, pv_component)],
        label='Battery Component', color='green', alpha=0.5)

# Add price indicators
ax2 = plt.gca().twinx()
ax2.plot(hours, electricity_prices, 'r--', label='Energy Price', linewidth=2)
ax2.set_ylabel('Price (€/kWh)', color='r')
ax2.tick_params(axis='y', labelcolor='r')

plt.xlabel('Hour')
plt.ylabel('Cost (€)')
plt.title('Final Cost Equation Components with Time-of-Use Prices')
plt.legend(loc='upper left')
plt.grid(True, alpha=0.3)

# Add total cost annotation
plt.text(0.02, 0.95, f'Total Cost F = {final_total_cost:.2f} €', 
         transform=plt.gca().transAxes, 
         bbox=dict(facecolor='white', alpha=0.8))

plt.tight_layout()
plt.show()

print("\nFinal Cost Equation Breakdown with Time-of-Use Prices:")
print("F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t + Pdischarge,t)) * price(t)]")
for t in range(24):
    print(f"\nHour {t} (Price: {electricity_prices[t]:.2f} €/kWh):")
    components = f"({power_vars[t].value():.2f} + {pv_power_vars[t].value():.2f} + ({charge_vars[t].value():.2f} + {discharge_vars[t].value():.2f})) * {electricity_prices[t]:.2f}"
    result = final_hourly_costs[t]
    print(f"  [{components}] = {result:.2f} €")

print(f"\nTotal F = {final_total_cost:.2f} €")

# Export calculations to Excel and Word
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create Word document
doc = Document()
doc.add_heading('Hybrid System Analysis Report', 0)

# Add introduction
doc.add_paragraph('This report presents the analysis of a hybrid power system consisting of wind turbine, PV system, and battery storage.')

# Add Wind Turbine section
doc.add_heading('Wind Turbine Analysis', level=1)
doc.add_paragraph(f'Rated Power: {Pr} kW')
doc.add_paragraph(f'Cut-in Speed: {vci} m/s')
doc.add_paragraph(f'Cut-out Speed: {vco} m/s')
doc.add_paragraph(f'Nominal Speed: {vN} m/s')

# Add PV System section
doc.add_heading('PV System Analysis', level=1)
doc.add_paragraph(f'Nominal Power: {Pn} kW')
doc.add_paragraph(f'Solar Irradiance Factor: {beta}')
doc.add_paragraph(f'Temperature Coefficient: {k}')

# Add Battery System section
doc.add_heading('Battery System Analysis', level=1)
doc.add_paragraph(f'Battery Capacity: {Ee} kWh')
doc.add_paragraph(f'Charging Efficiency: {eta_s}')
doc.add_paragraph(f'Discharging Efficiency: {eta_d}')
doc.add_paragraph(f'Maximum Power: {P_max} kW')

# Add hourly results table
doc.add_heading('Hourly Analysis', level=1)
table = doc.add_table(rows=1, cols=6)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
header_cells[0].text = 'Hour'
header_cells[1].text = 'Wind Power (kW)'
header_cells[2].text = 'PV Power (kW)'
header_cells[3].text = 'Battery Charge (kW)'
header_cells[4].text = 'Battery Discharge (kW)'
header_cells[5].text = 'Total Cost (€)'

for t in range(24):
    row_cells = table.add_row().cells
    row_cells[0].text = str(t)
    row_cells[1].text = f"{power_vars[t].value():.2f}"
    row_cells[2].text = f"{pv_power_vars[t].value():.2f}"
    row_cells[3].text = f"{charge_vars[t].value():.2f}"
    row_cells[4].text = f"{discharge_vars[t].value():.2f}"
    row_cells[5].text = f"{final_hourly_costs[t]:.2f}"

# Add final cost equation
doc.add_heading('Final Cost Analysis', level=1)
doc.add_paragraph('F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t - Pdischarge,t)) * price(t)]')
doc.add_paragraph(f'Final Total Cost: {final_total_cost:.2f} €').bold = True

# Add Calculation Sequence section
doc.add_heading('Calculation Sequence', level=1)

# Wind Turbine Calculations
doc.add_heading('Wind Turbine Power Calculation', level=2)
doc.add_paragraph('Wind Power Equation:')
doc.add_paragraph('If v < vci or v > vco: Power = 0')
doc.add_paragraph(f'If {vci} ≤ v ≤ {vN}: Power = {Pr} * ((v - {vci}) / ({vN} - {vci}))')
doc.add_paragraph(f'If {vN} < v ≤ {vco}: Power = {Pr}')

# Add wind calculation table
wind_table = doc.add_table(rows=1, cols=3)
wind_table.style = 'Table Grid'
wind_header = wind_table.rows[0].cells
wind_header[0].text = 'Hour'
wind_header[1].text = 'Wind Speed (m/s)'
wind_header[2].text = 'Power Output (kW)'

for t in range(24):
    row = wind_table.add_row().cells
    v = wind_speeds[t]
    row[0].text = str(t)
    row[1].text = f"{v:.2f}"
    if v < vci or v > vco:
        row[2].text = "P = 0 (Outside operating range)"
    elif vci <= v <= vN:
        row[2].text = f"P = {Pr} * ({v:.2f} - {vci}) / ({vN} - {vci})"
    else:
        row[2].text = f"P = {Pr} (Rated power)"
    row[2].text = f"{power_vars[t].value():.2f}"  # Fixed the format specifier

# PV System Calculations
doc.add_heading('PV System Power Calculation', level=2)
doc.add_paragraph('PV Power Equation:')
doc.add_paragraph(f'Power = {Pn} * {beta} * (G/1000) * (1 + {k} * (Tc - {Ts}))')
doc.add_paragraph('Where:')
doc.add_paragraph(f'G = Solar Irradiance (W/m²)\nTc = Cell Temperature (°C)')

# Add PV calculation table
pv_table = doc.add_table(rows=1, cols=4)
pv_table.style = 'Table Grid'
pv_header = pv_table.rows[0].cells
pv_header[0].text = 'Hour'
pv_header[1].text = 'Irradiance (W/m²)'
pv_header[2].text = 'Temperature (°C)'
pv_header[3].text = 'Power Output (kW)'

for t in range(24):
    row = pv_table.add_row().cells
    G = solar_irradiance[t]
    Tc = temperature[t]
    row[0].text = str(t)
    row[1].text = f"{G:.2f}"
    row[2].text = f"{Tc:.2f}"
    row[3].text = f"{pv_power_vars[t].value():.2f}"

# Battery Operation
doc.add_heading('Battery Operation Analysis', level=2)
doc.add_paragraph('Battery Constraints:')
doc.add_paragraph(f'State of Charge (SoC) limits: {SoC_min:.2f} kWh to {SoC_max:.2f} kWh')
doc.add_paragraph(f'Maximum Power: ±{P_max} kW')
doc.add_paragraph(f'Charging Efficiency: {eta_s*100}%')
doc.add_paragraph(f'Discharging Efficiency: {eta_d*100}%')

# Add battery operation table
bat_table = doc.add_table(rows=1, cols=4)
bat_table.style = 'Table Grid'
bat_header = bat_table.rows[0].cells
bat_header[0].text = 'Hour'
bat_header[1].text = 'Charging (kW)'
bat_header[2].text = 'Discharging (kW)'
bat_header[3].text = 'SoC (kWh)'

for t in range(24):
    row = bat_table.add_row().cells
    row[0].text = str(t)
    row[1].text = f"{charge_vars[t].value():.2f}"
    row[2].text = f"{discharge_vars[t].value():.2f}"
    row[3].text = f"{soc_vars[t].value():.2f}"

# Final Cost Calculation
doc.add_heading('Final Cost Calculation (F)', level=2)
doc.add_paragraph('Cost Equation:')
doc.add_paragraph('F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t - Pdischarge,t)) * price(t)]')

# Add detailed cost calculation table
cost_table = doc.add_table(rows=1, cols=6)
cost_table.style = 'Table Grid'
cost_header = cost_table.rows[0].cells
cost_header[0].text = 'Hour'
cost_header[1].text = 'Price (€/kWh)'
cost_header[2].text = 'Wind Cost (€)'
cost_header[3].text = 'PV Cost (€)'
cost_header[4].text = 'Battery Cost (€)'
cost_header[5].text = 'Total (€)'

for t in range(24):
    row = cost_table.add_row().cells
    row[0].text = str(t)
    row[1].text = f"{electricity_prices[t]:.4f}"
    row[2].text = f"{wind_component[t]:.2f}"
    row[3].text = f"{pv_component[t]:.2f}"
    row[4].text = f"{battery_component[t]:.2f}"
    row[5].text = f"{final_hourly_costs[t]:.2f}"

# Add final total
doc.add_paragraph(f'Final Total Cost F = {final_total_cost:.2f} €', style='Heading 2')

# Get the user's Documents directory and create file paths
import os
from pathlib import Path
documents_path = str(Path.home() / "Documents")

# Update file paths
word_file = os.path.join(documents_path, 'hybrid_system_report.docx')
excel_file = os.path.join(documents_path, 'hybrid_system_calculations.xlsx')

# Create a dictionary with all the data
data = {
    'Hour': hours,
    'Wind Speed (m/s)': wind_speeds,
    'Wind Power (kW)': [power_vars[t].value() for t in range(24)],
    'Solar Irradiance (W/m²)': solar_irradiance,
    'PV Power (kW)': [pv_power_vars[t].value() for t in range(24)],
    'Battery Charging (kW)': [charge_vars[t].value() for t in range(24)],
    'Battery Discharging (kW)': [discharge_vars[t].value() for t in range(24)],
    'Battery SoC (kWh)': [soc_vars[t].value() for t in range(24)],
    'Electricity Price (€/kWh)': electricity_prices,
    'Wind Cost (€)': wind_component,
    'PV Cost (€)': pv_component,
    'Battery Cost (€)': battery_component,
    'Total Hourly Cost (€)': final_hourly_costs
}

# Create DataFrame
df = pd.DataFrame(data)

# Add summary row
summary = pd.DataFrame({
    'Hour': ['Total'],
    'Wind Speed (m/s)': [sum(wind_speeds)],
    'Wind Power (kW)': [sum(power_vars[t].value() for t in range(24))],
    'Solar Irradiance (W/m²)': [sum(solar_irradiance)],
    'PV Power (kW)': [sum(pv_power_vars[t].value() for t in range(24))],
    'Battery Charging (kW)': [sum(charge_vars[t].value() for t in range(24))],
    'Battery Discharging (kW)': [sum(discharge_vars[t].value() for t in range(24))],
    'Battery SoC (kWh)': [soc_vars[23].value()],  # Final SoC
    'Electricity Price (€/kWh)': [sum(electricity_prices)/24],  # Average price
    'Wind Cost (€)': [sum(wind_component)],
    'PV Cost (€)': [sum(pv_component)],
    'Battery Cost (€)': [sum(battery_component)],
    'Total Hourly Cost (€)': [final_total_cost]
})

# Combine main data with summary
df = pd.concat([df, summary], ignore_index=True)

# Save Word document and Excel file
try:
    doc.save(word_file)
    print(f"\nReport exported to {word_file}")
except Exception as e:
    print(f"Error saving Word document: {e}")

# Export to Excel
try:
    with pd.ExcelWriter(excel_file, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Calculations', index=False)
        
        # Get workbook and worksheet objects
        workbook = writer.book
        worksheet = writer.sheets['Calculations']
        
        # Add formats
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#D7E4BC'
        })
        
        number_format = workbook.add_format({'num_format': '0.00'})
        
        # Apply formats
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            
        # Set column widths
        worksheet.set_column(0, len(df.columns)-1, 15)
        
        # Add final cost formula
        worksheet.write(len(df)+2, 0, 'Final Cost Equation:')
        worksheet.write(len(df)+3, 0, 'F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t - Pdischarge,t)) * price(t)]')
        worksheet.write(len(df)+4, 0, f'Final Total Cost F = {final_total_cost:.2f} €')
    
    print(f"\nCalculations exported to {excel_file}")
except Exception as e:
    print(f"Error saving Excel file: {e}")

# Create Mathematical Calculations Document
math_doc = Document()
math_doc.add_heading('Mathematical Calculations Report', 0)

# Wind Turbine Mathematical Analysis
math_doc.add_heading('Wind Turbine Mathematical Analysis', level=1)
math_doc.add_paragraph('Wind Power Function:')
math_doc.add_paragraph('P(v) = 0,                     if v < 3 m/s or v > 25 m/s')
math_doc.add_paragraph(f'P(v) = {Pr} * (v - {vci}) / ({vN} - {vci}),  if {vci} ≤ v ≤ {vN}')
math_doc.add_paragraph(f'P(v) = {Pr},                    if {vN} < v ≤ {vco}')

# Add wind power calculations table
math_wind_table = math_doc.add_table(rows=1, cols=4)
math_wind_table.style = 'Table Grid'
header_cells = math_wind_table.rows[0].cells
header_cells[0].text = 'Hour'
header_cells[1].text = 'Wind Speed (m/s)'
header_cells[2].text = 'Calculation'
header_cells[3].text = 'Power Output (kW)'

for t in range(24):
    row_cells = math_wind_table.add_row().cells
    v = wind_speeds[t]
    row_cells[0].text = str(t)
    row_cells[1].text = f"{v:.2f}"
    if v < vci or v > vco:
        row_cells[2].text = "P = 0 (Outside operating range)"
    elif vci <= v <= vN:
        row_cells[2].text = f"P = {Pr} * ({v:.2f} - {vci}) / ({vN} - {vci})"
    else:
        row_cells[2].text = f"P = {Pr} (Rated power)"
    row_cells[3].text = f"{power_vars[t].value():.2f}"  # Fixed the format specifier

# PV System Mathematical Analysis
math_doc.add_heading('PV System Mathematical Analysis', level=1)
math_doc.add_paragraph('PV Power Equation:')
math_doc.add_paragraph(f'P = {Pn} * {beta} * (G/1000) * (1 + {k} * (Tc - {Ts}))')

# Add PV calculations table
math_pv_table = math_doc.add_table(rows=1, cols=5)
math_pv_table.style = 'Table Grid'
header_cells = math_pv_table.rows[0].cells
header_cells[0].text = 'Hour'
header_cells[1].text = 'Irradiance (W/m²)'
header_cells[2].text = 'Temperature (°C)'
header_cells[3].text = 'Calculation'
header_cells[4].text = 'Power Output (kW)'

for t in range(24):
    row_cells = math_pv_table.add_row().cells
    G = solar_irradiance[t]
    Tc = temperature[t]
    row_cells[0].text = str(t)
    row_cells[1].text = f"{G:.2f}"
    row_cells[2].text = f"{Tc:.2f}"
    row_cells[3].text = f"P = {Pn}*{beta}*({G:.0f}/1000)*(1+{k}*({Tc:.1f}-{Ts}))"
    row_cells[4].text = f"{pv_power_vars[t].value():.2f}"

# Battery System Mathematical Analysis
math_doc.add_heading('Battery System Mathematical Analysis', level=1)
math_doc.add_paragraph('Battery State of Charge Evolution:')
math_doc.add_paragraph(f'SoC(t+1) = SoC(t) + (η_c * P_c(t) - P_d(t)/η_d) * Δt')
math_doc.add_paragraph(f'where: η_c = {eta_s}, η_d = {eta_d}')

# Add battery calculations table
math_bat_table = math_doc.add_table(rows=1, cols=5)
math_bat_table.style = 'Table Grid'
header_cells = math_bat_table.rows[0].cells
header_cells[0].text = 'Hour'
header_cells[1].text = 'Charging (kW)'
header_cells[2].text = 'Discharging (kW)'
header_cells[3].text = 'SoC Evolution'
header_cells[4].text = 'SoC (kWh)'

for t in range(24):
    row_cells = math_bat_table.add_row().cells
    row_cells[0].text = str(t)
    row_cells[1].text = f"{charge_vars[t].value():.2f}"
    row_cells[2].text = f"{discharge_vars[t].value():.2f}"
    row_cells[3].text = f"SoC({t+1}) = SoC({t}) + ({eta_s}*{charge_vars[t].value():.2f} - {discharge_vars[t].value():.2f}/{eta_d})"
    row_cells[4].text = f"{soc_vars[t].value():.2f}"

# Final Cost Function Mathematical Analysis
math_doc.add_heading('Cost Function Mathematical Analysis', level=1)
math_doc.add_paragraph('Final Cost Function F:')
math_doc.add_paragraph('F = Σ(t=1 to 24) [(Pwind,t + Ppv,t + (Pcharge,t - Pdischarge,t)) * price(t)]')

# Add cost calculations table
math_cost_table = math_doc.add_table(rows=1, cols=6)
math_cost_table.style = 'Table Grid'
header_cells = math_cost_table.rows[0].cells
header_cells[0].text = 'Hour'
header_cells[1].text = 'Components'
header_cells[2].text = 'Price (€/kWh)'
header_cells[3].text = 'Calculation'
header_cells[4].text = 'Result (€)'
header_cells[5].text = 'Cumulative (€)'

cumulative_cost = 0
for t in range(24):
    row_cells = math_cost_table.add_row().cells
    wind_term = power_vars[t].value()
    pv_term = pv_power_vars[t].value()
    charge_term = charge_vars[t].value()
    discharge_term = discharge_vars[t].value()
    price = electricity_prices[t]
    hourly_cost = final_hourly_costs[t]
    cumulative_cost += hourly_cost
    
    row_cells[0].text = str(t)
    row_cells[1].text = f"Wind: {wind_term:.2f}\nPV: {pv_term:.2f}\nCharge: {charge_term:.2f}\nDischarge: {discharge_term:.2f}"
    row_cells[2].text = f"{price:.4f}"
    row_cells[3].text = f"({wind_term:.2f} + {pv_term:.2f} + ({charge_term:.2f} - {discharge_term:.2f})) * {price:.4f}"
    row_cells[4].text = f"{hourly_cost:.2f}"
    row_cells[5].text = f"{cumulative_cost:.2f}"

# Add final result
math_doc.add_paragraph(f'Final Total Cost F = {final_total_cost:.2f} €', style='Heading 2')

# Save mathematical calculations document
math_file = os.path.join(documents_path, 'mathematical_calculations.docx')
try:
    math_doc.save(math_file)
    print(f"\nMathematical calculations exported to {math_file}")
except Exception as e:
    print(f"Error saving mathematical calculations: {e}")
